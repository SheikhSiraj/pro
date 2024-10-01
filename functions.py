from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
import base64
import streamlit as st

def set_page_style():
    st.markdown(
        """
        <style>
            /* Define light and dark mode variables */
            :root {
                --bg-color-light: #f7f7f9;
                --bg-color-dark: #343a40;
                --text-color-light: #343a40;
                --text-color-dark: #f7f7f9;
                --input-bg-light: rgba(255, 255, 255, 0.9);
                --input-bg-dark: rgba(40, 40, 40, 0.9);
                --input-border-light: #007bff;
                --input-border-dark: #007bff;
                --input-text-color-light: #000000;
                --input-text-color-dark: #ffffff;
            }

            /* Use body background and text colors based on theme */
            body {
                background-color: var(--bg-color-light);  /* Light mode background */
                color: var(--text-color-light);  /* Light mode text color */
                font-family: 'Arial', sans-serif;
                transition: background-color 0.3s, color 0.3s;
            }

            body[data-theme='dark'] {
                background-color: var(--bg-color-dark);  /* Dark mode background */
                color: var(--text-color-dark);  /* Dark mode text color */
            }

            /* Button styling */
            .stButton button {
                background: linear-gradient(90deg, #007bff, #00c6ff);  /* Gradient background */
                color: white;
                padding: 12px 24px;
                border: none;
                border-radius: 30px;
                cursor: pointer;
                font-size: 16px;
                transition: background 0.3s, transform 0.3s;
                box-shadow: 0 4px 10px rgba(0, 123, 255, 0.3);
            }
            .stButton button:hover {
                background: linear-gradient(90deg, #0056b3, #00a1e7);
                transform: translateY(-2px);
            }

            /* Input field styling */
            .stTextInput, .stTextArea, .stSelectbox, .stNumberInput, .stDateInput {
                padding: 12px;
                border-radius: 20px;
                border: 2px solid var(--input-border-light);
                margin-bottom: 15px;
                box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
                transition: border-color 0.3s, background-color 0.3s;
                background: var(--input-bg-light);
                color: var(--input-text-color-light);
            }

            body[data-theme='dark'] .stTextInput, 
            body[data-theme='dark'] .stTextArea, 
            body[data-theme='dark'] .stSelectbox, 
            body[data-theme='dark'] .stNumberInput, 
            body[data-theme='dark'] .stDateInput {
                background: var(--input-bg-dark);
                color: var(--input-text-color-dark);
                border-color: var(--input-border-dark);
            }

            .stTextInput input, .stTextArea textarea, .stNumberInput input, .stDateInput input {
                border: none;
                width: 100%;
                font-size: 16px;
                padding: 12px;
                outline: none;
                color: inherit;  /* Inherit color from parent for better visibility */
            }

            /* Skills section (Multiselect) styling */
            .stMultiselect div[role="combobox"] {
                padding: 12px;
                border: 2px solid var(--input-border-light);
                border-radius: 20px;
                transition: border-color 0.3s, background-color 0.3s;
                background: var(--input-bg-light);
                color: var(--input-text-color-light);
            }

            body[data-theme='dark'] .stMultiselect div[role="combobox"] {
                background: var(--input-bg-dark);
                color: var(--input-text-color-dark);
                border-color: var(--input-border-dark);
            }

            /* Bold labels for input fields */
            .stSelectbox label, .stTextInput label, .stTextArea label, .stNumberInput label, .stDateInput label, .stMultiselect label {
                font-weight: bold;
                font-size: 15px;
                margin-bottom: 5px;
            }

            /* Main content box styling */
            .main-content {
                padding: 30px;
                background-color: white;
                border-radius: 15px;
                box-shadow: 0 4px 20px rgba(0, 0, 0, 0.1);
                max-width: 800px;
                margin: auto;
            }

            body[data-theme='dark'] .main-content {
                background-color: #1e1e1e;  /* Dark mode content background */
            }
        </style>
        """,
        unsafe_allow_html=True
    )



# Function to generate the DOCX file
def generate_docx(details, font_name):
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)

    # Add border
    sectPr = sections[0]._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(ns.qn('w:val'), 'single')
        border.set(ns.qn('w:sz'), '4')
        border.set(ns.qn('w:space'), '24')
        border.set(ns.qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)

    # Add Name, Email, Phone, LinkedIn, and Website
    name = doc.add_heading(level=1)
    run = name.add_run(details['name'].upper())  # Convert name to uppercase
    run.font.size = Pt(24)
    run.font.name = font_name
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_info = doc.add_paragraph()
    contact_text = f"Email: {details['email']} | Phone: {details['phone']}"
    if details.get('linkedin'):
        contact_text += f" | LinkedIn: {details['linkedin']}"
    if details.get('website'):
        contact_text += f" | Website: {details['website']}"
    run = contact_info.add_run(contact_text)
    run.font.name = font_name
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Summary', level=2)
    summary_paragraph = doc.add_paragraph(details['summary'])
    for run in summary_paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Add Experience if available
    if details['experience']:
        doc.add_heading('Experience', level=2)
        for exp in details['experience']:
            if exp['job_title'] and exp['company']:
                exp_paragraph = doc.add_paragraph()
                exp_paragraph.add_run(f"{exp['job_title']} - {exp['company']}").bold = True
                exp_paragraph.add_run(f"\n{exp['start_date']} to {exp['end_date']}")
                exp_paragraph.add_run(f"\n{exp['responsibilities']}")

    # Add Education
    doc.add_heading('Education', level=2)
    for edu in details['education']:
        if edu['degree'] and edu['institution']:
            edu_paragraph = doc.add_paragraph()
            edu_paragraph.add_run(f"{edu['degree']} - {edu['institution']}").bold = True
            edu_paragraph.add_run(f"\nGraduation Date: {edu['graduation_date']}")

    # Add Skills
    doc.add_heading('Skills', level=2)
    skill_paragraph = doc.add_paragraph()
    skill_paragraph.add_run(", ".join(details['skills']))

    # Add Projects
    if details['projects']:
        doc.add_heading('Projects', level=2)
        for proj in details['projects']:
            if proj['title']:
                proj_paragraph = doc.add_paragraph()
                proj_paragraph.add_run(f"{proj['title']}").bold = True
                proj_paragraph.add_run(f"\nRole: {proj['role']}")
                proj_paragraph.add_run(f"\n{proj['description']}")

    # Add Awards
    if details['awards']:
        doc.add_heading('Awards', level=2)
        for award in details['awards']:
            if award['title']:
                award_paragraph = doc.add_paragraph()
                award_paragraph.add_run(f"{award['title']}").bold = True
                award_paragraph.add_run(f"\nReason: {award['reason']}")

    return doc

# Function to create the download link for the generated resume
def create_download_link(val, filename):
    b64 = base64.b64encode(val).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download your resume</a>'

# Adding custom CSS for background image
def set_background_image(image_path):
    with open(image_path, "rb") as image_file:
        encoded_image = base64.b64encode(image_file.read()).decode('utf-8')
    
    st.markdown(
        f"""
        <style>
        .stApp {{
            background: url('data:image/png;base64,{encoded_image}');
            background-size: cover;
            background-position: center;
            background-attachment: fixed;
            height: 100vh;  /* Ensure full viewport height */
        }}
        </style>
        """,
        unsafe_allow_html=True
    )
